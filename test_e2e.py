# -*- coding: utf-8 -*-
"""
Pruebas End-to-End (E2E) del generador de reportes de lubricantes.

METODOLOGIA: exclusivamente E2E con archivos REALES. No se usa ningun
mock / monkeypatch / stub. Cada prueba ejecuta el flujo completo real:

    datos -> Excel real -> lectura real -> render real con docxtpl
          -> archivo .docx real en disco -> verificacion abriendo el .docx

Se ejecuta con pytest:        pytest test_e2e.py -v
o directamente con Python:    python test_e2e.py
"""
import base64
import datetime
import os
import shutil
import tempfile

import openpyxl
from docx import Document

import report_model as m
import prepare_template
import history

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(HERE, "plantilla_reporte.docx")
ORIGINAL = os.path.join(HERE, "reporte_original.docx")

# PNG real de 1x1 pixel (datos binarios reales, no es un mock).
_PNG_1x1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAAC0lEQVR42mNk"
    "YAAAAAYAAjCB0C8AAAAASUVORK5CYII=")


def _workdir():
    d = tempfile.mkdtemp(prefix="lubes_e2e_")
    return d


def _ensure_template():
    """Garantiza que exista la plantilla docxtpl real (la genera si falta)."""
    if os.path.isfile(TEMPLATE):
        return
    if not os.path.isfile(ORIGINAL):
        raise RuntimeError(
            "Falta 'plantilla_reporte.docx' y 'reporte_original.docx'. "
            "Ejecute prepare_template.py con el reporte Word original.")
    import sys
    sys.argv = ["prepare_template.py", ORIGINAL]
    prepare_template.main()


def _all_text(doc):
    """Concatena todo el texto de parrafos y tablas del documento."""
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# --------------------------------------------------------------------------
# E2E 1: ida y vuelta real por Excel
# --------------------------------------------------------------------------

def test_e2e_excel_roundtrip():
    work = _workdir()
    try:
        xlsx = os.path.join(work, "datos.xlsx")
        original = m.default_data()
        m.save_excel(original, xlsx)
        assert os.path.isfile(xlsx), "El Excel no se escribio en disco"

        loaded = m.load_excel(xlsx)
        assert loaded["period_full"] == original["period_full"]
        assert len(loaded["consolidated"]) == len(original["consolidated"])
        # Los valores numericos sobreviven la ida y vuelta.
        assert loaded["consolidated"][0]["opening"] == \
            original["consolidated"][0]["opening"]
        assert loaded["deliveries"]["S4CX30"][0]["docket"] == \
            original["deliveries"]["S4CX30"][0]["docket"]
        print("OK  test_e2e_excel_roundtrip")
    finally:
        shutil.rmtree(work, ignore_errors=True)


# --------------------------------------------------------------------------
# E2E 2: flujo completo Excel -> reporte Word y verificacion del .docx
# --------------------------------------------------------------------------

def test_e2e_full_generation_from_excel():
    _ensure_template()
    work = _workdir()
    try:
        xlsx = os.path.join(work, "datos.xlsx")
        out = os.path.join(work, "reporte.docx")

        m.save_excel(m.default_data(), xlsx)
        data = m.load_excel(xlsx)
        result = m.generate_report(data, {}, TEMPLATE, out)

        assert os.path.isfile(result), "No se genero el archivo .docx"
        doc = Document(result)
        text = _all_text(doc)

        # No deben quedar etiquetas Jinja sin renderizar.
        assert "{{" not in text and "{%" not in text, \
            "Quedaron etiquetas de plantilla sin procesar"

        # La tabla consolidada tiene encabezado + 4 productos.
        cons = doc.tables[0]
        assert len(cons.rows) == 5
        sites = [cons.rows[i].cells[0].text for i in range(1, 5)]
        assert sites == m.PRODUCT_KEYS

        # Calculo derivado real: calc_stock = opening + deliveries - transactions.
        # S4CX30: 16576 + 11417 - 8718 = 19275
        assert cons.rows[1].cells[4].text == "19.275"

        # Tabla de entregas del producto 1: 2 tickets + total = 4 filas.
        assert len(doc.tables[2].rows) == 4
        assert doc.tables[2].rows[3].cells[0].text == "Grand Total"

        # Textos clave presentes.
        assert "Issued product SPIRAX S4CX30." in text
        assert "Dispensing to Equipment: 6.732 L." in text
        print("OK  test_e2e_full_generation_from_excel")
    finally:
        shutil.rmtree(work, ignore_errors=True)


# --------------------------------------------------------------------------
# E2E 3: datos modificados se reflejan en el reporte real
# --------------------------------------------------------------------------

def test_e2e_custom_data_reflected():
    _ensure_template()
    work = _workdir()
    try:
        out = os.path.join(work, "reporte.docx")
        data = m.default_data()
        data["period_full"] = "June 1st to 8th, 2026"
        data["consolidated"][0]["opening"] = 99999
        data["deliveries"]["S4CX30"] = [
            {"date": "02/06/2026", "volume": 1000, "docket": 900},
        ]
        m.generate_report(data, {}, TEMPLATE, out)

        text = _all_text(Document(out))
        assert "June 1st to 8th, 2026" in text
        assert "99.999" in text          # opening modificado
        assert "02/06/2026" in text      # nuevo ticket de entrega
        # Variance del ticket = 900 - 1000 = -100
        assert "-100" in text
        print("OK  test_e2e_custom_data_reflected")
    finally:
        shutil.rmtree(work, ignore_errors=True)


# --------------------------------------------------------------------------
# E2E 4: insercion real de imagenes en el reporte
# --------------------------------------------------------------------------

def test_e2e_image_insertion():
    _ensure_template()
    work = _workdir()
    try:
        png = os.path.join(work, "figura.png")
        with open(png, "wb") as fh:
            fh.write(_PNG_1x1)

        out_no = os.path.join(work, "sin_imagenes.docx")
        out_yes = os.path.join(work, "con_imagenes.docx")

        m.generate_report(m.default_data(), {}, TEMPLATE, out_no)
        # fig9, fig13 y fig_tasks NO se autogeneran: deben sumar 3 imagenes.
        images = {"fig9": png, "fig13": png, "fig_tasks": png}
        m.generate_report(m.default_data(), images, TEMPLATE, out_yes)

        shapes_no = len(Document(out_no).inline_shapes)
        shapes_yes = len(Document(out_yes).inline_shapes)
        assert shapes_yes == shapes_no + 3, \
            "Se esperaban 3 imagenes nuevas insertadas (%d vs %d)" % (
                shapes_yes, shapes_no)
        print("OK  test_e2e_image_insertion")
    finally:
        shutil.rmtree(work, ignore_errors=True)


# --------------------------------------------------------------------------
# E2E 5: el porcentaje de variance se calcula correctamente
# --------------------------------------------------------------------------

def test_e2e_variance_calculation():
    _ensure_template()
    work = _workdir()
    try:
        out = os.path.join(work, "reporte.docx")
        data = m.default_data()
        # 15W40: variance = closing - (opening + deliveries - transactions)
        #               = 20571 - (23843 + 0 - 3317) = 20571 - 20526 = 45
        # % = 45 / 3317 * 100 = 1,36%
        m.generate_report(data, {}, TEMPLATE, out)
        cons = Document(out).tables[0]
        assert cons.rows[2].cells[0].text == "15W40"
        assert cons.rows[2].cells[7].text == "45"      # variance
        assert cons.rows[2].cells[8].text == "1,36%"   # porcentaje
        print("OK  test_e2e_variance_calculation")
    finally:
        shutil.rmtree(work, ignore_errors=True)


# --------------------------------------------------------------------------
# E2E 6: extraccion por fecha desde un Excel historico real
# --------------------------------------------------------------------------

def _build_history_xlsx(path):
    """Crea un Excel historico REAL con el formato de las hojas 'Recon ...'.
    Dos semanas; en la segunda, el producto S5CFDM60 no tiene fila de datos.
    """
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    headers = ["Date", "Site", "Opening Stock", "Deliveries", "Transactions",
               "Calculated stock", "Net Stock change", "Closing stock",
               "Variance", "%", None, "To equipment", "Other Dispenses",
               "Transfers"]
    # producto -> filas (date, site, opening, deliveries, transactions,
    #                     closing, to_equipment, other, transfers)
    sheets = {
        "Recon Spirax S4CX30": [
            # S4CX30 tiene varios tanques: solo debe tomarse "Tank 21".
            (datetime.datetime(2026, 5, 4), "Tank 21", 1000, 500, 300, 1200,
             250, 0, 50),
            (datetime.datetime(2026, 5, 4), "Tank 2", 99999, 99999, 99999,
             99999, 99999, 0, 0),
            (datetime.datetime(2026, 5, 11), "Tank 21", 1200, 600, 400, 1380,
             350, 0, 50),
            (datetime.datetime(2026, 5, 11), "Tank 2", 88888, 88888, 88888,
             88888, 88888, 0, 0),
        ],
        "Recon 15W40": [
            (datetime.datetime(2026, 5, 4), "Tank 6", 2000, 0, 800, 1200,
             800, 0, 0),
            (datetime.datetime(2026, 5, 11), "Tank 6", 1200, 900, 500, 1600,
             500, 0, 0),
        ],
        "Recon Spirax S5CFDM60": [
            (datetime.datetime(2026, 5, 4), "Tank 4", 400, 300, 100, 600,
             100, 0, 0),
            # 2026-05-11 SIN datos: solo la fecha, resto vacio.
        ],
        "Recon Tellus S3M46": [
            (datetime.datetime(2026, 5, 4), "Tank 9", 300, 200, 150, 350,
             0, 0, 150),
            (datetime.datetime(2026, 5, 11), "Tank 9", 350, 250, 180, 420,
             0, 0, 180),
        ],
    }
    for name, rows in sheets.items():
        ws = wb.create_sheet(name)
        ws.append(headers)
        ws.append([None] * 14)
        for r in rows:
            date, site, op, deliv, trans, close, eq, ot, tr = r
            ws.append([date, site, op, deliv, trans, op + deliv - trans,
                       close - op, close, 0, 0, None, eq, ot, tr])
        if name == "Recon Spirax S5CFDM60":
            ws.append([datetime.datetime(2026, 5, 11)] + [None] * 13)

    # Hoja de tickets de entrega (formato real "delivery_transaction_...").
    dt = wb.create_sheet("delivery_transaction_213_202312")
    dt.append(["Product", "Collected At", "Tank", "Docket Number", "Supplier",
               "Confirmed", "Type", "Volume", "Docket Volume", "Variance", "%"])
    dt.append(["Spirax S4CX30", datetime.datetime(2026, 5, 7, 9, 0), "Tank 2",
               "D-IN", None, "Yes", "Gauged", 5000, 4800, -200, 0])
    dt.append(["Spirax S4CX30", datetime.datetime(2026, 5, 2, 9, 0), "Tank 2",
               "D-OUT", None, "Yes", "Gauged", 9999, 9999, 0, 0])
    dt.append(["15W40", datetime.datetime(2026, 5, 9, 9, 0), "Tank 6",
               "D-NO", None, "No", "Gauged", 3000, 3000, 0, 0])

    # Hojas "WeeklyVariance ..." (tendencia de variance por producto).
    wv_sheets = {
        "WeeklyVariance  Spirax S4CX30": "S4CX30",
        "WeeklyVariance Spirax S4CX10W": "S4CX10W",
        "WeeklyVariance  15W40": "15W40",
        "WeeklyVariance  Spirax S5CFDM60": "S5CFDM60",
        "Weekly Variance  Tellus S3M46": "Tellus S3M46",
    }
    for name in wv_sheets:
        ws = wb.create_sheet(name)
        ws.append(["Weekly variance change", None, None])
        ws.append(["Date", "Delivery %", "Recon %"])
        ws.append([datetime.datetime(2026, 4, 27), 0, -2.5])
        ws.append([datetime.datetime(2026, 5, 4), 0, 3.1])
        ws.append([datetime.datetime(2026, 5, 11), 0, -1.8])
    wb.save(path)


def _build_stock_trend_csv(path):
    """Crea un CSV real con el formato 'stock_trend_...' del FMS."""
    import csv
    header = ["", "",
              "Tank WS - Spirax S4CX30 - Tank 2",
              "Tank WS - Spirax S4CX10W - Tank 1",
              "Tank WS - Rimula R4X15W40 - Tank 6",
              "Tank WS - Spirax S5CFDM60 - Tank 4",
              "Tank WS - Hydraulic Fluid - Tellus S3M46 - Tank 9"]
    with open(path, "w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh)
        w.writerow(header)
        w.writerow(["Safe Fill Capacity for Product", "Litres",
                    9000, 14000, 30000, 8000, 8000])
        base = datetime.datetime(2026, 5, 4)
        for i in range(20):
            ts = (base + datetime.timedelta(hours=8 * i)
                  ).strftime("%Y-%m-%d %H:%M:%S -0300")
            w.writerow([ts, "Litres", 2800 + i * 30, 13800 - i * 50,
                        16500 + i * 40, 6000 - i * 10, 3800 + i * 5])


def test_e2e_history_date_lookup():
    work = _workdir()
    try:
        xlsx = os.path.join(work, "historico.xlsx")
        _build_history_xlsx(xlsx)

        store = history.HistoryStore()
        store.load(xlsx)
        assert store.is_loaded()

        dates = store.available_dates()
        assert datetime.date(2026, 5, 4) in dates
        assert datetime.date(2026, 5, 11) in dates

        # Semana con datos completos para S4CX30.
        rec = store.lookup("S4CX30", datetime.date(2026, 5, 11))
        assert rec is not None
        assert rec.opening == 1200 and rec.closing == 1380
        assert rec.transfers == 50
        # Debe usar SOLO "Tank 21"; el decoy "Tank 2" (88888) se ignora.
        assert rec.tanks == ["Tank 21"], "S4CX30 debe usar solo Tank 21"
        assert rec.opening != 88888 and rec.transactions != 88888

        # Fecha SIN informacion registrada -> debe devolver None.
        missing = store.lookup("S5CFDM60", datetime.date(2026, 5, 11))
        assert missing is None, "Debio reportar fecha no registrada"
        print("OK  test_e2e_history_date_lookup")
    finally:
        shutil.rmtree(work, ignore_errors=True)


def test_e2e_history_extraction_to_report():
    _ensure_template()
    work = _workdir()
    try:
        xlsx = os.path.join(work, "historico.xlsx")
        out = os.path.join(work, "reporte.docx")
        _build_history_xlsx(xlsx)

        store = history.HistoryStore()
        store.load(xlsx)
        data = m.default_data()

        # Extraer la semana 2026-05-11 (S5CFDM60 no esta registrado).
        found, missing = history.apply_week_to_data(
            data, store, datetime.date(2026, 5, 11))
        assert set(found) == {"S4CX30", "15W40", "Tellus S3M46"}
        assert missing == ["S5CFDM60"]

        # Los datos extraidos llegan al reporte Word real.
        m.generate_report(data, {}, TEMPLATE, out)
        cons = Document(out).tables[0]
        # S4CX30: opening extraido = 1200 -> formateado "1.200".
        assert cons.rows[1].cells[0].text == "S4CX30"
        assert cons.rows[1].cells[1].text == "1.200"
        print("OK  test_e2e_history_extraction_to_report")
    finally:
        shutil.rmtree(work, ignore_errors=True)


def test_e2e_delivery_tickets_by_week():
    """Los tickets de entrega se filtran por la semana seleccionada y solo
    se incluyen los confirmados."""
    _ensure_template()
    work = _workdir()
    try:
        xlsx = os.path.join(work, "historico.xlsx")
        out = os.path.join(work, "reporte.docx")
        _build_history_xlsx(xlsx)

        store = history.HistoryStore()
        store.load(xlsx)
        data = m.default_data()
        history.apply_week_to_data(data, store, datetime.date(2026, 5, 11))

        # Solo el ticket dentro de la semana (07/05) y confirmado.
        s4 = data["deliveries"]["S4CX30"]
        assert len(s4) == 1, "Debio quedar solo el ticket de la semana"
        assert s4[0]["date"] == "07/05/2026"
        assert s4[0]["volume"] == 5000
        # El ticket de 15W40 estaba 'Confirmed=No' -> se descarta.
        assert data["deliveries"]["15W40"] == []

        m.generate_report(data, {}, TEMPLATE, out)
        text = _all_text(Document(out))
        assert "07/05/2026" in text
        assert "9999" not in text and "9.999" not in text   # ticket fuera de semana
        print("OK  test_e2e_delivery_tickets_by_week")
    finally:
        shutil.rmtree(work, ignore_errors=True)


def test_e2e_auto_charts_generated():
    """Las Figuras 1-8 (barras de entregas y donas de transacciones) se
    construyen solas: aparecen en el reporte aunque no se adjunte imagen."""
    _ensure_template()
    work = _workdir()
    try:
        out = os.path.join(work, "reporte.docx")
        # Sin adjuntar ninguna imagen: igual deben insertarse 8 graficos.
        m.generate_report(m.default_data(), {}, TEMPLATE, out)
        shapes = len(Document(out).inline_shapes)
        assert shapes == 8, \
            "Se esperaban 8 graficos autogenerados (fig1-8), hay %d" % shapes
        print("OK  test_e2e_auto_charts_generated")
    finally:
        shutil.rmtree(work, ignore_errors=True)


def test_e2e_auto_delivery_narrative():
    """La narrativa de la figura de entregas cuenta los tickets reales."""
    _ensure_template()
    work = _workdir()
    try:
        out = os.path.join(work, "reporte.docx")
        data = m.default_data()
        # S4CX30 con 2 tickets -> "Two confirmed deliveries".
        data["deliveries"]["S4CX30"] = [
            {"date": "05/05/2026", "volume": 1000, "docket": 900},
            {"date": "07/05/2026", "volume": 2000, "docket": 1800},
        ]
        # 15W40 sin tickets -> "Zero confirmed delivery".
        data["deliveries"]["15W40"] = []
        m.generate_report(data, {}, TEMPLATE, out)
        text = _all_text(Document(out))
        assert "Two confirmed deliveries" in text
        assert "Zero confirmed delivery" in text
        print("OK  test_e2e_auto_delivery_narrative")
    finally:
        shutil.rmtree(work, ignore_errors=True)


def test_e2e_cover_date_and_period():
    """La fecha de portada y el periodo se aplican de forma consistente y la
    fecha vieja del template ya no aparece."""
    _ensure_template()
    work = _workdir()
    try:
        out = os.path.join(work, "reporte.docx")
        data = m.default_data()
        data["period_full"] = "April 30th to May 6th, 2026"
        data["cover_date"] = "May 6th, 2026"
        m.generate_report(data, {}, TEMPLATE, out)

        import zipfile
        xml = zipfile.ZipFile(out).read(
            "word/document.xml").decode("utf-8", "ignore")
        alltext = "".join(__import__("re").findall(
            r"<w:t[^>]*>([^<]*)</w:t>", xml))
        assert "May 6th, 2026" in alltext, "fecha de portada no aplicada"
        assert "May 18th" not in alltext, "quedo la fecha vieja de la portada"
        assert "April 30th to May 6th, 2026" in alltext

        # history.period_for deriva periodo y portada de la fecha elegida.
        xlsx = os.path.join(work, "historico.xlsx")
        _build_history_xlsx(xlsx)
        store = history.HistoryStore()
        store.load(xlsx)
        period, cover = store.period_for(datetime.date(2026, 5, 11))
        assert cover == "May 11th, 2026", cover
        assert period == "May 4th to May 11th, 2026", period
        print("OK  test_e2e_cover_date_and_period")
    finally:
        shutil.rmtree(work, ignore_errors=True)


def test_e2e_recon_trends_from_weeklyvariance():
    """Las tendencias de variance se leen de las hojas WeeklyVariance."""
    work = _workdir()
    try:
        xlsx = os.path.join(work, "historico.xlsx")
        _build_history_xlsx(xlsx)
        store = history.HistoryStore()
        store.load(xlsx)
        trend = store.recon_trend("15W40", datetime.date(2026, 5, 11))
        assert len(trend) == 3, "Se esperaban 3 puntos de tendencia"
        assert trend[-1] == (datetime.date(2026, 5, 11), -1.8)
        # Un punto posterior a la fecha pedida queda excluido.
        trend2 = store.recon_trend("15W40", datetime.date(2026, 5, 4))
        assert len(trend2) == 2
        print("OK  test_e2e_recon_trends_from_weeklyvariance")
    finally:
        shutil.rmtree(work, ignore_errors=True)


def test_e2e_all_18_figures_autogenerated():
    """Con el Excel historico + el CSV de tanques, las 18 figuras se generan
    solas (fig_tasks queda manual)."""
    _ensure_template()
    work = _workdir()
    try:
        xlsx = os.path.join(work, "historico.xlsx")
        csvp = os.path.join(work, "stock_trend.csv")
        out = os.path.join(work, "reporte.docx")
        _build_history_xlsx(xlsx)
        _build_stock_trend_csv(csvp)

        store = history.HistoryStore()
        store.load(xlsx)
        data = m.default_data()
        history.apply_week_to_data(data, store, datetime.date(2026, 5, 11))

        import stock_trend
        trends, safe = stock_trend.load_tank_trends(csvp)
        data["tank_trends"] = trends
        data["tank_safe_fill"] = safe

        m.generate_report(data, {}, TEMPLATE, out)
        shapes = len(Document(out).inline_shapes)
        assert shapes == 18, \
            "Se esperaban 18 figuras autogeneradas, hay %d" % shapes
        print("OK  test_e2e_all_18_figures_autogenerated")
    finally:
        shutil.rmtree(work, ignore_errors=True)


if __name__ == "__main__":
    tests = [
        test_e2e_excel_roundtrip,
        test_e2e_full_generation_from_excel,
        test_e2e_custom_data_reflected,
        test_e2e_image_insertion,
        test_e2e_variance_calculation,
        test_e2e_history_date_lookup,
        test_e2e_history_extraction_to_report,
        test_e2e_delivery_tickets_by_week,
        test_e2e_auto_charts_generated,
        test_e2e_auto_delivery_narrative,
        test_e2e_cover_date_and_period,
        test_e2e_recon_trends_from_weeklyvariance,
        test_e2e_all_18_figures_autogenerated,
    ]
    failed = 0
    for test in tests:
        try:
            test()
        except Exception as exc:
            failed += 1
            print("FALLO  %s: %s" % (test.__name__, exc))
    print("\n%d/%d pruebas E2E superadas." % (len(tests) - failed, len(tests)))
    raise SystemExit(1 if failed else 0)
