# -*- coding: utf-8 -*-
"""
Convierte el reporte Word original en una PLANTILLA docxtpl con etiquetas
Jinja ({{ ... }} y {%tr ... %}).

Se ejecuta UNA sola vez (o cada vez que cambie el diseno base del Word).
El usuario final NO necesita tocar Word manualmente: este script inserta
todas las etiquetas en los lugares correctos.

Uso:
    python prepare_template.py "ruta/al/reporte_original.docx"
    (si se omite la ruta, busca 'reporte_original.docx' en esta carpeta)

Genera:  plantilla_reporte.docx
"""
import copy
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

DEFAULT_FONT = "Arial"

# Detecta frases de periodo tipo "May 11th to 18th, 2026" o
# "May 11th to May 18th, 2026" en cualquier parrafo (leyendas de tablas y
# figuras). Todas se reemplazan por la etiqueta {{ period_full }} para que
# coincidan con el periodo elegido por el usuario.
PERIOD_RE = re.compile(
    r"[A-Z][a-z]+\s+\d{1,2}(?:st|nd|rd|th)?\s+to\s+"
    r"(?:[A-Z][a-z]+\s+)?\d{1,2}(?:st|nd|rd|th)?\s*,?\s*\d{4}")

# Detecta una fecha suelta tipo "May 18th, 2026" (la fecha de la portada,
# que vive dentro de cuadros de texto).
STANDALONE_DATE_RE = re.compile(
    r"[A-Z][a-z]+\s+\d{1,2}(?:st|nd|rd|th)?\s*,?\s*\d{4}")

OUTPUT_NAME = "plantilla_reporte.docx"

# Indices (basados en el reporte original entregado) de los parrafos que
# contienen las figuras, en orden de documento.
FIGURE_PARAGRAPHS = {
    60: "fig1", 79: "fig2", 102: "fig3", 125: "fig4",
    142: "fig5", 165: "fig6", 198: "fig7", 232: "fig8",
    265: "fig9", 268: "fig10", 280: "fig11", 283: "fig12",
    297: "fig13", 300: "fig14", 317: "fig15", 320: "fig16",
    334: "fig17", 337: "fig18", 379: "fig_tasks",
}

# Parrafos de texto que se reemplazan completos por una etiqueta.
# Las etiquetas {{r ... }} indican RichText (Arial + negritas inline).
TEXT_PARAGRAPHS = {
    42: "From {{ period_full }}",
    49: "{{r p1_recon_sentence }}",
    58: "{{r p1_delivery_narrative }}",
    64: "{{r p2_recon_sentence }}",
    77: "{{r p2_delivery_narrative }}",
    87: "{{r p3_recon_sentence }}",
    100: "{{r p3_delivery_narrative }}",
    110: "{{r p4_recon_sentence }}",
    123: "{{r p4_delivery_narrative }}",
    # Bloque dispensing producto 1
    140: "{{r p1_issued }}",
    144: "{{r p1_disp_total }}",
    146: "{{r p1_disp_equipment_line }}",
    147: "{{r p1_disp_other_line }}",
    148: "{{r p1_disp_transfers_line }}",
    # Bloque dispensing producto 2
    163: "{{r p2_issued }}",
    167: "{{r p2_disp_total }}",
    169: "{{r p2_disp_equipment_line }}",
    170: "{{r p2_disp_other_line }}",
    171: "{{r p2_disp_transfers_line }}",
    # Bloque dispensing producto 3
    196: "{{r p3_issued }}",
    200: "{{r p3_disp_total }}",
    202: "{{r p3_disp_equipment_line }}",
    203: "{{r p3_disp_other_line }}",
    204: "{{r p3_disp_transfers_line }}",
    # Bloque dispensing producto 4
    230: "{{r p4_issued }}",
    234: "{{r p4_disp_total }}",
    236: "{{r p4_disp_equipment_line }}",
    237: "{{r p4_disp_other_line }}",
    238: "{{r p4_disp_transfers_line }}",
}

# Grupos de parrafos de "Considerations" -> se convierten en un bucle Jinja.
CONSIDERATION_GROUPS = {
    "p1_cons": [52, 53],
    "p2_cons": [68, 69, 70, 71],
    "p3_cons": [91, 92, 93, 94],
    "p4_cons": [114, 115, 116, 117],
}


def set_paragraph_text(paragraph, text):
    """Borra el contenido del parrafo y lo reemplaza por `text` en Arial
    (preservando el formato del primer run cuando existe)."""
    _rewrite_keep_format(paragraph, text)


def delete_paragraph(paragraph):
    el = paragraph._p
    el.getparent().remove(el)


def _rewrite_keep_format(paragraph, new_text):
    """Reemplaza todo el contenido del parrafo por `new_text`, conservando
    el formato del primer run (y forzando fuente Arial cuando el original
    no tenia fuente explicita, para igualar el reporte hecho a mano)."""
    runs = paragraph.runs
    bold = italic = underline = None
    name = None
    size = None
    if runs:
        r0 = runs[0]
        bold, italic, underline = r0.bold, r0.italic, r0.underline
        name = r0.font.name
        size = r0.font.size
    for child in list(paragraph._p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            paragraph._p.remove(child)
    run = paragraph.add_run(new_text)
    run.bold, run.italic, run.underline = bold, italic, underline
    run.font.name = name or DEFAULT_FONT
    if size is not None:
        run.font.size = size


def retag_period(paragraph):
    """Si el parrafo contiene una fecha de periodo, la sustituye por la
    etiqueta {{ period_full }}."""
    text = paragraph.text
    if not PERIOD_RE.search(text) or "{{" in text or "{%" in text:
        return False
    _rewrite_keep_format(paragraph, PERIOD_RE.sub("{{ period_full }}", text))
    return True


def retag_textbox_dates(doc):
    """Etiqueta las fechas que viven dentro de cuadros de texto (la portada).
    Una fecha suelta -> {{ cover_date }}; un rango -> {{ period_full }}."""
    count = 0
    for txbx in doc.element.iter(qn("w:txbxContent")):
        for p_el in txbx.iter(qn("w:p")):
            para = Paragraph(p_el, doc)
            text = para.text
            if not text.strip() or "{{" in text or "{%" in text:
                continue
            if PERIOD_RE.search(text):
                _rewrite_keep_format(
                    para, PERIOD_RE.sub("{{ period_full }}", text))
                count += 1
            elif STANDALONE_DATE_RE.search(text):
                _rewrite_keep_format(
                    para, STANDALONE_DATE_RE.sub("{{ cover_date }}", text))
                count += 1
    return count


def set_cell(cell, text):
    cell.text = text


def _clone_control_row(data_row, control_text):
    """Crea un <w:tr> clon que contiene solo una directiva de control Jinja.
    docxtpl reemplaza la fila completa por la directiva, asi que el resto
    del contenido del clon es irrelevante (se vacia por prolijidad)."""
    new_tr = copy.deepcopy(data_row._tr)
    return new_tr


def _insert_loop_rows(table, data_row, for_text, end_text):
    """Inserta una fila de control antes y otra despues de data_row."""
    for_tr = copy.deepcopy(data_row._tr)
    end_tr = copy.deepcopy(data_row._tr)
    data_row._tr.addprevious(for_tr)
    data_row._tr.addnext(end_tr)
    # Vaciar las celdas de las filas de control y poner la directiva.
    for tr, text in ((for_tr, for_text), (end_tr, end_text)):
        cells = tr.findall(".//" + qn("w:tc"))
        for ci, tc in enumerate(cells):
            for para in tc.findall(qn("w:p")):
                for child in list(para):
                    if child.tag in (qn("w:r"), qn("w:hyperlink")):
                        para.remove(child)
            if ci == 0:
                run = para.makeelement(qn("w:r"), {})
                wt = para.makeelement(qn("w:t"), {})
                wt.text = text
                run.append(wt)
                para.append(run)


def tag_consolidated_table(table):
    """Tabla 1: tabla consolidada -> bucle {%tr%} con filas de control."""
    fields = ["site", "opening", "deliveries", "transactions",
              "calc_stock", "net_change", "closing", "variance", "pct"]
    data_row = table.rows[1]
    _insert_loop_rows(table, data_row,
                      "{%tr for r in consolidated %}", "{%tr endfor %}")
    for i, field in enumerate(fields):
        set_cell(data_row.cells[i], "{{ r.%s }}" % field)
    # Eliminar las filas de datos originales sobrantes (otros productos).
    for extra in list(table.rows[4:]):
        table._tbl.remove(extra._tr)


def tag_stock_table(table, prefix):
    """Tablas de stock por producto (una sola fila de datos)."""
    row = table.rows[1]
    fields = ["site", "opening", "deliveries", "transactions",
              "calc_stock", "net_change", "closing", "variance", "pct"]
    for i, field in enumerate(fields):
        set_cell(row.cells[i], "{{ %s_stock.%s }}" % (prefix, field))


def tag_delivery_table(table, prefix):
    """Tablas de entregas por producto: bucle de filas + fila Grand Total."""
    fields = ["date", "volume", "docket", "variance", "pct"]
    data_row = table.rows[1]
    _insert_loop_rows(table, data_row,
                      "{%%tr for d in %s_deliveries %%}" % prefix,
                      "{%tr endfor %}")
    for i, field in enumerate(fields):
        set_cell(data_row.cells[i], "{{ d.%s }}" % field)
    # Filas de datos intermedias (entre la fila bucle y el Grand Total).
    for extra in list(table.rows[4:-1]):
        table._tbl.remove(extra._tr)
    total_row = table.rows[-1]
    set_cell(total_row.cells[0], "Grand Total")
    set_cell(total_row.cells[1], "{{ %s_dtot.volume }}" % prefix)
    set_cell(total_row.cells[2], "{{ %s_dtot.docket }}" % prefix)
    set_cell(total_row.cells[3], "{{ %s_dtot.variance }}" % prefix)
    set_cell(total_row.cells[4], "{{ %s_dtot.pct }}" % prefix)


def main():
    if len(sys.argv) > 1:
        src = sys.argv[1]
    else:
        src = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "reporte_original.docx")
    if not os.path.isfile(src):
        print("ERROR: no se encontro el reporte original: %s" % src)
        print("Uso: python prepare_template.py \"ruta/al/reporte_original.docx\"")
        sys.exit(1)

    doc = Document(src)
    paragraphs = list(doc.paragraphs)

    # 1) Capturar referencias ANTES de mutar (los indices cambian al borrar).
    text_targets = {idx: paragraphs[idx] for idx in TEXT_PARAGRAPHS}
    figure_targets = {idx: paragraphs[idx] for idx in FIGURE_PARAGRAPHS}
    cons_targets = {name: [paragraphs[i] for i in idxs]
                    for name, idxs in CONSIDERATION_GROUPS.items()}

    # 2) Reemplazar parrafos de texto.
    for idx, tag in TEXT_PARAGRAPHS.items():
        set_paragraph_text(text_targets[idx], tag)

    # 3) Reemplazar figuras por etiquetas de imagen.
    for idx, fig in FIGURE_PARAGRAPHS.items():
        set_paragraph_text(figure_targets[idx], "{{ %s }}" % fig)

    # 4) Convertir cada grupo de consideraciones en un bucle Jinja.
    for name, paras in cons_targets.items():
        first = paras[0]
        first.insert_paragraph_before("{%%p for c in %s %%}" % name,
                                      style=first.style)
        set_paragraph_text(first, "{{ c }}")
        for middle in paras[1:-1]:
            delete_paragraph(middle)
        set_paragraph_text(paras[-1], "{%p endfor %}")

    # 5) Etiquetar las 9 tablas.
    tables = doc.tables
    tag_consolidated_table(tables[0])
    tag_stock_table(tables[1], "p1")
    tag_delivery_table(tables[2], "p1")
    tag_stock_table(tables[3], "p2")
    tag_delivery_table(tables[4], "p2")
    tag_stock_table(tables[5], "p3")
    tag_delivery_table(tables[6], "p3")
    tag_stock_table(tables[7], "p4")
    tag_delivery_table(tables[8], "p4")

    # 6) Unificar TODAS las fechas de periodo (leyendas de tablas y figuras)
    #    con la etiqueta {{ period_full }}.
    retagged = sum(retag_period(p) for p in doc.paragraphs)
    print("Fechas de periodo unificadas en %d leyendas." % retagged)

    # 7) Etiquetar la fecha de la portada (vive en cuadros de texto).
    cover = retag_textbox_dates(doc)
    print("Fechas de portada/cuadros de texto etiquetadas: %d." % cover)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), OUTPUT_NAME)
    doc.save(out)
    print("Plantilla generada correctamente: %s" % out)


if __name__ == "__main__":
    main()
